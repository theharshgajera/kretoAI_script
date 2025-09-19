from flask import Flask, render_template, request, jsonify, session
from flask_cors import CORS
from dotenv import load_dotenv
import os
import google.generativeai as genai
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound, VideoUnavailable
import time
import re
import uuid
import json
from datetime import datetime
import logging
from urllib.parse import urlparse, parse_qs
import threading
from collections import defaultdict

# Document processing imports
import PyPDF2
import docx
from werkzeug.utils import secure_filename
import tempfile
import fitz  # PyMuPDF for better PDF handling
from pathlib import Path

# Load environment variables
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "your-secret-key-here")
CORS(app)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# File upload configuration
UPLOAD_FOLDER = os.path.join(tempfile.gettempdir(), 'script_generator_uploads')
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}
MAX_CONTENT_LENGTH = 50 * 1024 * 1024  # 50MB max file size

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Ensure upload folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Enhanced in-memory storage
user_data = defaultdict(lambda: {
    'folders': {},
    'processing_status': {},
    'analysis_cache': {},
    'chat_sessions': {},
    'current_script': None,
    'insights_cache': {},
    'documents': {},  # New: Store document data
    'document_insights': {}  # New: Store document analysis
})

class DocumentProcessor:
    """Advanced document processing with multiple format support"""
    
    def __init__(self):
        self.max_chars = 100000  # Maximum characters to process per document
    
    def allowed_file(self, filename):
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
    
    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF using multiple methods for reliability"""
        text = ""
        
        try:
            # Try PyMuPDF first (better for complex PDFs)
            doc = fitz.open(file_path)
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text += page.get_text()
            doc.close()
            
            if len(text.strip()) > 50:  # If we got good text
                return text
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")
        
        try:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {e}")
            return None
        
        return text if len(text.strip()) > 50 else None
    
    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n'.join(text)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return None
    
    def extract_text_from_doc(self, file_path):
        """Extract text from legacy DOC files (basic support)"""
        try:
            # Try using python-docx (limited DOC support)
            doc = docx.Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            return '\n'.join(text)
        except Exception as e:
            logger.warning(f"DOC extraction failed, file might need conversion: {e}")
            return None
    
    def extract_text_from_txt(self, file_path):
        """Extract text from TXT files with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"TXT extraction failed: {e}")
                return None
        
        return None
    
    def process_document(self, file_path, filename):
        """Process document and extract text with metadata"""
        file_ext = filename.rsplit('.', 1)[1].lower()
        
        extraction_methods = {
            'pdf': self.extract_text_from_pdf,
            'docx': self.extract_text_from_docx,
            'doc': self.extract_text_from_doc,
            'txt': self.extract_text_from_txt
        }
        
        extract_method = extraction_methods.get(file_ext)
        if not extract_method:
            return {"error": "Unsupported file format", "text": None, "stats": None}
        
        try:
            text = extract_method(file_path)
            
            if not text or len(text.strip()) < 50:
                return {"error": "Could not extract meaningful text from document", "text": None, "stats": None}
            
            # Truncate if too long
            if len(text) > self.max_chars:
                text = text[:self.max_chars] + "\n\n[Document truncated for processing...]"
            
            stats = self._calculate_document_stats(text, filename)
            
            return {
                "error": None,
                "text": text,
                "stats": stats,
                "filename": filename,
                "file_type": file_ext
            }
            
        except Exception as e:
            logger.error(f"Document processing error: {e}")
            return {"error": f"Error processing document: {str(e)}", "text": None, "stats": None}
    
    def _calculate_document_stats(self, text, filename):
        """Calculate document statistics"""
        if not text:
            return {
                'char_count': 0,
                'word_count': 0,
                'page_estimate': 0,
                'read_time': 0
            }
        
        char_count = len(text)
        word_count = len(text.split())
        page_estimate = max(1, word_count // 250)  # ~250 words per page
        read_time = max(1, word_count // 200)  # ~200 words per minute
        
        return {
            'char_count': char_count,
            'word_count': word_count,
            'page_estimate': page_estimate,
            'read_time': read_time,
            'filename': filename
        }

class VideoProcessor:
    """Existing video processor - keeping all your advanced logic"""
    def __init__(self):
        self.rate_limit_delay = 2
        self.last_api_call = 0
    
    def extract_video_id(self, youtube_url):
        patterns = [
            r'(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/)([^&\n?#]+)',
            r'youtube\.com\/watch\?.*v=([^&\n?#]+)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, youtube_url)
            if match:
                return match.group(1)
        return None
    
    def validate_youtube_url(self, url):
        youtube_domains = ['youtube.com', 'youtu.be', 'www.youtube.com']
        try:
            parsed_url = urlparse(url)
            return any(domain in parsed_url.netloc for domain in youtube_domains)
        except:
            return False
    
    def rate_limit_wait(self):
        current_time = time.time()
        time_since_last_call = current_time - self.last_api_call
        if time_since_last_call < self.rate_limit_delay:
            sleep_time = self.rate_limit_delay - time_since_last_call
            time.sleep(sleep_time)
        self.last_api_call = time.time()
    
    def extract_transcript_details(self, youtube_video_url, max_retries=3, retry_delay=2):
        """Your existing advanced transcript extraction logic"""
        self.rate_limit_wait()
        
        video_id = self.extract_video_id(youtube_video_url)
        if not video_id:
            return {"error": "Invalid YouTube URL format", "transcript": None, "stats": None}
        
        for attempt in range(max_retries):
            try:
                if attempt > 0:
                    wait_time = retry_delay * (attempt + 1)
                    logger.info(f"Retrying transcript extraction in {wait_time} seconds... (Attempt {attempt + 1}/{max_retries})")
                    time.sleep(wait_time)
                
                ytt_api = YouTubeTranscriptApi()
                
                try:
                    fetched_transcript = ytt_api.fetch(video_id, languages=['en', 'en-US', 'en-GB'])
                    transcript_text = " ".join([snippet.text for snippet in fetched_transcript])
                    
                    if len(transcript_text.strip()) >= 50:
                        stats = self._calculate_transcript_stats(transcript_text)
                        logger.info(f"Direct fetch successful - {stats['char_count']} characters")
                        return {"error": None, "transcript": transcript_text, "stats": stats}
                        
                except NoTranscriptFound:
                    logger.info("Direct fetch failed, trying transcript list method")
                    pass
                
                try:
                    transcript_list = ytt_api.list(video_id)
                    
                    transcript = None
                    try:
                        transcript = transcript_list.find_transcript(['en', 'en-US', 'en-GB'])
                        logger.info("Found English transcript")
                    except NoTranscriptFound:
                        try:
                            transcript = transcript_list.find_manually_created_transcript(['en', 'en-US', 'en-GB'])
                            logger.info("Found manually created English transcript")
                        except NoTranscriptFound:
                            try:
                                transcript = transcript_list.find_generated_transcript(['en', 'en-US', 'en-GB'])
                                logger.info("Found auto-generated English transcript")
                            except NoTranscriptFound:
                                available_transcripts = list(transcript_list)
                                if available_transcripts:
                                    transcript = available_transcripts[0]
                                    logger.info(f"Using first available transcript: {transcript.language}")
                                else:
                                    return {"error": "No transcripts available for this video", "transcript": None, "stats": None}
                    
                    if not transcript:
                        return {"error": "No suitable transcript found", "transcript": None, "stats": None}
                    
                    fetched_transcript = transcript.fetch()
                    transcript_text = " ".join([snippet.text for snippet in fetched_transcript])
                    
                    if len(transcript_text.strip()) < 50:
                        return {"error": "Transcript too short or incomplete", "transcript": None, "stats": None}
                    
                    stats = self._calculate_transcript_stats(transcript_text)
                    logger.info(f"Transcript extraction successful - {stats['char_count']} characters, {stats['word_count']} words")
                    return {"error": None, "transcript": transcript_text, "stats": stats}
                    
                except Exception as inner_e:
                    logger.error(f"Inner exception during transcript list processing: {str(inner_e)}")
                    if attempt == max_retries - 1:
                        return {"error": f"Could not access transcript list - {str(inner_e)}", "transcript": None, "stats": None}
                    continue

            except VideoUnavailable:
                return {"error": "Video is unavailable, private, or doesn't exist", "transcript": None, "stats": None}
            except TranscriptsDisabled:
                return {"error": "Transcripts are disabled for this video", "transcript": None, "stats": None}
            except Exception as e:
                error_msg = str(e).lower()
                logger.error(f"Attempt {attempt + 1} failed: {str(e)}")
                
                if "quota" in error_msg or "rate" in error_msg or "429" in error_msg:
                    if attempt < max_retries - 1:
                        logger.warning(f"Rate limit hit. Waiting {retry_delay * (attempt + 2)} seconds before retry...")
                        continue
                    else:
                        return {"error": "API rate limit exceeded. Please try again in a few minutes", "transcript": None, "stats": None}
                elif "403" in error_msg or "forbidden" in error_msg:
                    return {"error": "Access forbidden. Video might be private or restricted", "transcript": None, "stats": None}
                elif "404" in error_msg:
                    return {"error": "Video not found. Please check the URL", "transcript": None, "stats": None}
                elif "blocked" in error_msg or "ipblocked" in error_msg:
                    return {"error": "IP address blocked by YouTube. Try using a VPN or proxy", "transcript": None, "stats": None}
                elif attempt == max_retries - 1:
                    return {"error": f"Error fetching transcript after {max_retries} attempts: {str(e)}", "transcript": None, "stats": None}
        
        return {"error": "Failed to fetch transcript after multiple attempts", "transcript": None, "stats": None}
    
    def _calculate_transcript_stats(self, transcript_text):
        if not transcript_text:
            return {
                'char_count': 0,
                'word_count': 0,
                'estimated_duration': 0,
                'estimated_read_time': 0
            }
        
        char_count = len(transcript_text)
        word_count = len(transcript_text.split())
        estimated_duration = max(1, word_count // 150)
        estimated_read_time = max(1, word_count // 200)
        
        return {
            'char_count': char_count,
            'word_count': word_count,
            'estimated_duration': estimated_duration,
            'estimated_read_time': estimated_read_time
        }

class EnhancedScriptGenerator:
    """Your existing advanced script generator with document integration"""
    
    def __init__(self):
        # Your existing advanced prompts
        self.style_analysis_prompt = """
        You are an expert YouTube content analyst. Analyze the following transcripts from the creator's personal videos to create a comprehensive style profile.

        Focus on identifying:

        **VOICE & TONE CHARACTERISTICS:**
        - Speaking style (conversational, formal, energetic, calm, etc.)
        - Emotional tone and energy levels
        - Use of humor, sarcasm, or specific personality traits
        - Level of enthusiasm and passion
        - Pacing and rhythm patterns
        
        **LANGUAGE PATTERNS:**
        - Vocabulary complexity and word choices
        - Sentence structure preferences (short/long, simple/complex)
        - Catchphrases, repeated expressions, or signature sayings
        - Use of technical jargon vs. simple explanations
        - Storytelling approach and narrative style
        - Transition phrases and connection words
        
        **CONTENT STRUCTURE & FLOW:**
        - How they introduce topics and hook viewers
        - Transition techniques between sections
        - How they build up to main points
        - Conclusion and call-to-action styles
        - Use of examples, analogies, and explanations
        - Information presentation patterns
        
        **ENGAGEMENT TECHNIQUES:**
        - How they ask questions to audience
        - Interactive elements and audience engagement
        - Use of personal stories and experiences
        - How they handle complex topics
        - Teaching and explanation methodology
        - Retention strategies used
        
        **UNIQUE CREATOR CHARACTERISTICS:**
        - What makes this creator distinctive
        - Their unique perspective or angle
        - Personal brand elements
        - Values and beliefs that come through
        - Specific expertise areas and how they showcase them
        - Content themes and recurring topics

        **KEY INSIGHTS FOR SCRIPT GENERATION:**
        - Most effective hooks and openings used
        - Common content structures that work well
        - Signature explanations or teaching methods
        - Audience connection techniques
        - Call-to-action patterns

        Provide a detailed, actionable style profile that captures the creator's authentic voice for script generation.

        **Creator's Personal Video Transcripts:**
        """
        
        self.inspiration_analysis_prompt = """
        You are an expert content strategist and topic analyst. Analyze these inspiration video transcripts to extract valuable insights and identify key topics with detailed breakdowns.

        Extract and organize:

        **CORE TOPICS & DETAILED INSIGHTS:**
        - Main subject matters with specific subtopics
        - Key points and arguments presented
        - Data, statistics, and factual claims
        - Expert opinions and industry insights
        - Trending discussions and current debates
        - Evergreen vs. timely content themes
        
        **CONTENT IDEAS & CREATIVE ANGLES:**
        - Unique perspectives and fresh takes on topics
        - Creative approaches to common subjects
        - Unexplored angles or missing viewpoints
        - Potential spin-offs and related topics
        - Cross-topic connection opportunities
        - Controversial or debate-worthy points
        
        **STORYTELLING & PRESENTATION TECHNIQUES:**
        - Narrative structures and story arcs used
        - How complex topics are simplified
        - Types of examples and case studies used
        - Visual or conceptual metaphors
        - Emotional appeals and connection methods
        - Pacing and information delivery patterns
        
        **VALUABLE INSIGHTS & ACTIONABLE INFORMATION:**
        - Specific tips, tricks, and how-to steps
        - Common problems and detailed solution approaches
        - Industry best practices mentioned
        - Tools, resources, and recommendations
        - Success stories and failure case studies
        - Expert advice and professional insights
        
        **TOPIC-SPECIFIC MAIN POINTS BREAKDOWN:**
        For each major topic discussed, provide:
        - Core concept explanation
        - Key supporting arguments
        - Practical applications mentioned
        - Common misconceptions addressed
        - Advanced concepts introduced
        - Related subtopics worth exploring
        
        **CONTENT GAPS & OPPORTUNITIES:**
        - Topics that could be expanded upon
        - Alternative viewpoints not covered
        - Beginner vs. advanced treatment opportunities
        - Updated information or fresh perspectives needed
        - Underexplored subtopics with potential

        Provide a comprehensive analysis that captures both the content insights and the presentation methods for creating informed, original content.

        **Inspiration Video Transcripts:**
        """

        # NEW: Document analysis prompt
        self.document_analysis_prompt = """
        You are an expert content analyst specializing in document comprehension and insight extraction. Analyze the following document content to extract key insights, main points, and actionable information that can inform YouTube script generation.

        Focus on identifying:

        **CORE CONCEPTS & MAIN THEMES:**
        - Primary topics and subject areas covered
        - Key concepts and definitions
        - Central arguments and thesis points
        - Supporting evidence and data
        - Expert opinions and authoritative insights
        
        **ACTIONABLE INFORMATION:**
        - Step-by-step processes and procedures
        - Specific tips, strategies, and recommendations
        - Tools, resources, and methodologies mentioned
        - Best practices and proven approaches
        - Case studies and real-world examples
        
        **KNOWLEDGE STRUCTURE:**
        - Logical flow of information
        - How concepts build upon each other
        - Prerequisites and foundational knowledge needed
        - Advanced concepts and expert-level insights
        - Practical applications and implementations
        
        **CONTENT OPPORTUNITIES FOR VIDEO SCRIPTS:**
        - Main points that could become video topics
        - Complex concepts that need simplification
        - Practical demonstrations or tutorials possible
        - Controversial or debate-worthy points
        - Current vs. outdated information
        - Gaps that could be filled with additional research
        
        **AUDIENCE VALUE PROPOSITIONS:**
        - What viewers would learn or gain
        - Problems this content helps solve
        - Skills or knowledge they would acquire
        - Practical benefits and outcomes
        - Target audience level (beginner/intermediate/advanced)

        Extract the most valuable insights that could inform comprehensive, educational YouTube content creation.

        **Document Content:**
        """

        # Enhanced script generation with document integration
        self.enhanced_script_template = """
        You are an expert YouTube script writer creating a professional, engaging script based on comprehensive content analysis including the creator's style, topic insights, and document knowledge.

        **CREATOR'S AUTHENTIC STYLE PROFILE:**
        {style_profile}

        **TOPIC INSIGHTS FROM INSPIRATION CONTENT:**
        {inspiration_summary}

        **DOCUMENT KNOWLEDGE BASE:**
        {document_insights}

        **USER'S SPECIFIC REQUEST:**
        {user_prompt}

        **ENHANCED SCRIPT GENERATION INSTRUCTIONS:**

        Create a complete, production-ready YouTube script that:

        1. **MAINTAINS AUTHENTIC VOICE:** Use the creator's natural speaking style, vocabulary, and personality traits identified in the style profile.

        2. **INTEGRATES DOCUMENT KNOWLEDGE STRATEGICALLY:**
           - Use document insights as authoritative foundation
           - Incorporate specific data, facts, and expert knowledge
           - Reference key concepts and methodologies from documents
           - Build upon documented best practices and proven approaches
           
        3. **LEVERAGES INSPIRATION INSIGHTS:**
           - Include trending discussions and current debates
           - Use successful presentation techniques identified
           - Apply proven engagement strategies
           - Reference industry insights and expert opinions

        4. **FOLLOWS PROFESSIONAL STRUCTURE:**
           - **Hook (0-15 seconds):** Attention-grabbing opening with specific value promise
           - **Introduction (15-45 seconds):** Topic setup with authority establishment
           - **Main Content Sections:** Well-structured body with clear progression
           - **Document-Based Authority:** Weave in expert knowledge naturally
           - **Practical Applications:** Include actionable takeaways
           - **Conclusion:** Strong summary with clear next steps

        5. **ENSURES COMPREHENSIVE COVERAGE:**
           - Address the topic from multiple angles identified in documents
           - Include both foundational and advanced concepts appropriately
           - Provide practical examples and real-world applications
           - Reference credible sources and expert insights
           - Balance theory with actionable advice

        6. **MAINTAINS ENGAGEMENT:**
           - Use the creator's proven engagement techniques
           - Include questions, interactions, and retention hooks
           - Apply storytelling methods that resonate
           - Incorporate appropriate humor or personality elements

        **OUTPUT FORMAT:**
        
        # [COMPELLING VIDEO TITLE]
        
        ## CONTENT FOUNDATION
        **Document Authority:** [Key document insights being leveraged]
        **Topic Relevance:** [Why this matters now based on inspiration analysis]
        **Creator Angle:** [Unique perspective based on style profile]
        
        ## HOOK (0-15 seconds)
        [Attention-grabbing opening with authority and promise]
        **[Production Note: Tone, visual, and delivery guidance]**
        
        ## INTRODUCTION (15-45 seconds)  
        [Authority establishment with document-backed credibility]
        **[Expert Insight: Specific fact or data from documents]**
        
        ## MAIN CONTENT
        
        ### Section 1: [Title] (Timing: X:XX - X:XX)
        [Document-informed content with creator's authentic delivery]
        **[Authority Point: Specific expert knowledge from documents]**
        **[Actionable Takeaway: Practical application]**
        **[Production Note: Visual aids, emphasis cues]**
        
        ### Section 2: [Title] (Timing: X:XX - X:XX)
        [Continue with comprehensive, well-researched content]
        **[Expert Validation: Supporting evidence from documents]**
        **[Real-World Application: How viewers implement this]**
        
        [Continue for all main sections...]
        
        ## CONCLUSION (Last 30-60 seconds)
        [Strong summary with document-backed authority and clear next steps]
        **[Final Authority Statement: Key expert insight that reinforces value]**
        
        ---
        
        **PRODUCTION NOTES:**
        - Visual timeline and supporting materials needed
        - Key emphasis points for authority and credibility
        - Document references and source citations
        - Graphics, data visualization opportunities
        - Expert quote overlays or callouts
        
        **AUTHORITY & CREDIBILITY ELEMENTS:**
        - Document insights strategically integrated
        - Expert knowledge naturally woven throughout
        - Factual backing for all major claims
        - Credible source references where appropriate
        - Balance of accessible explanation with authoritative depth

        Generate the complete script now, ensuring it authentically matches the creator's style while delivering authoritative, document-informed content on the requested topic.
        """

        # Your existing chat modification prompt
        self.chat_modification_prompt = """
        You are an expert YouTube script editor working with a creator to refine their script. You have access to:

        **ORIGINAL SCRIPT:**
        {current_script}

        **CREATOR'S STYLE PROFILE:**
        {style_profile}

        **TOPIC INSIGHTS:**
        {topic_insights}

        **DOCUMENT KNOWLEDGE:**
        {document_insights}

        **MODIFICATION REQUEST:**
        {user_message}

        **INSTRUCTIONS:**
        Based on the creator's request, modify the script while:
        1. Maintaining their authentic voice and style
        2. Keeping document authority and expert knowledge intact
        3. Preserving key insights and valuable information
        4. Making targeted improvements based on the specific request
        5. Ensuring any new content fits naturally with existing flow

        **RESPONSE FORMAT:**
        **Modified Script:**
        [Provide the updated script or specific sections that were changed]

        **Changes Made:**
        - [Bullet point list of specific changes]
        - [Explanation of why these changes improve the script]
        - [Any suggestions for further improvements]

        Respond as if you're collaborating with the creator in a natural conversation.
        """
    
    def analyze_creator_style(self, personal_transcripts):
        """Your existing enhanced style analysis"""
        combined_transcripts = "\n\n---VIDEO SEPARATOR---\n\n".join(personal_transcripts)
        
        max_chars = 50000
        if len(combined_transcripts) > max_chars:
            chunk_size = max_chars // 3
            start_chunk = combined_transcripts[:chunk_size]
            middle_start = len(combined_transcripts) // 2 - chunk_size // 2
            middle_chunk = combined_transcripts[middle_start:middle_start + chunk_size]
            end_chunk = combined_transcripts[-chunk_size:]
            combined_transcripts = f"{start_chunk}\n\n[...CONTENT CONTINUES...]\n\n{middle_chunk}\n\n[...CONTENT CONTINUES...]\n\n{end_chunk}"
        
        try:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(
                self.style_analysis_prompt + combined_transcripts,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.1,
                    max_output_tokens=3000
                )
            )
            
            if response.text:
                return response.text
            else:
                return "Could not analyze creator style - empty response"
                
        except Exception as e:
            logger.error(f"Error analyzing creator style: {str(e)}")
            return f"Error analyzing creator style: {str(e)}"
    
    def analyze_inspiration_content(self, inspiration_transcripts):
        """Your existing enhanced inspiration analysis"""
        combined_transcripts = "\n\n---VIDEO SEPARATOR---\n\n".join(inspiration_transcripts)
        
        max_chars = 50000
        if len(combined_transcripts) > max_chars:
            chunk_size = max_chars // len(inspiration_transcripts) if len(inspiration_transcripts) > 1 else max_chars
            sampled_transcripts = []
            for transcript in inspiration_transcripts:
                if len(transcript) > chunk_size:
                    half_chunk = chunk_size // 2
                    sampled = transcript[:half_chunk] + "\n[...CONTENT CONTINUES...]\n" + transcript[-half_chunk:]
                    sampled_transcripts.append(sampled)
                else:
                    sampled_transcripts.append(transcript)
            combined_transcripts = "\n\n---VIDEO SEPARATOR---\n\n".join(sampled_transcripts)
        
        try:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(
                self.inspiration_analysis_prompt + combined_transcripts,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.2,
                    max_output_tokens=3000
                )
            )
            
            if response.text:
                return response.text
            else:
                return "Could not analyze inspiration content - empty response"
                
        except Exception as e:
            logger.error(f"Error analyzing inspiration content: {str(e)}")
            return f"Error analyzing inspiration content: {str(e)}"
    
    def analyze_documents(self, document_texts):
        """NEW: Analyze uploaded documents for insights"""
        if not document_texts:
            return "No documents provided for analysis."
        
        combined_documents = "\n\n---DOCUMENT SEPARATOR---\n\n".join(document_texts)
        
        # Limit document content for processing
        max_chars = 60000
        if len(combined_documents) > max_chars:
            # Take chunks from each document rather than truncating
            chunk_size = max_chars // len(document_texts) if len(document_texts) > 1 else max_chars
            sampled_docs = []
            for doc_text in document_texts:
                if len(doc_text) > chunk_size:
                    # Take beginning and end of each document
                    half_chunk = chunk_size // 2
                    sampled = doc_text[:half_chunk] + "\n[...DOCUMENT CONTINUES...]\n" + doc_text[-half_chunk:]
                    sampled_docs.append(sampled)
                else:
                    sampled_docs.append(doc_text)
            combined_documents = "\n\n---DOCUMENT SEPARATOR---\n\n".join(sampled_docs)
        
        try:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(
                self.document_analysis_prompt + combined_documents,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.2,
                    max_output_tokens=3500
                )
            )
            
            if response.text:
                return response.text
            else:
                return "Could not analyze document content - empty response"
                
        except Exception as e:
            logger.error(f"Error analyzing documents: {str(e)}")
            return f"Error analyzing documents: {str(e)}"
    
    def generate_enhanced_script(self, style_profile, inspiration_summary, document_insights, user_prompt):
        """Generate script with all available knowledge sources"""
        
        enhanced_prompt = self.enhanced_script_template.format(
            style_profile=style_profile,
            inspiration_summary=inspiration_summary,
            document_insights=document_insights,
            user_prompt=user_prompt
        )
        
        try:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(
                enhanced_prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.7,
                    max_output_tokens=4000
                )
            )
            
            if response.text:
                return response.text
            else:
                return "Error: Could not generate script - empty response"
                
        except Exception as e:
            logger.error(f"Error generating enhanced script: {str(e)}")
            return f"Error generating script: {str(e)}"

    def modify_script_chat(self, current_script, style_profile, topic_insights, document_insights, user_message):
        """Modify script with full context including documents"""
        
        chat_prompt = self.chat_modification_prompt.format(
            current_script=current_script,
            style_profile=style_profile,
            topic_insights=topic_insights,
            document_insights=document_insights,
            user_message=user_message
        )
        
        try:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(
                chat_prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.6,
                    max_output_tokens=3000
                )
            )
            
            if response.text:
                return response.text
            else:
                return "Could not modify script - empty response"
                
        except Exception as e:
            logger.error(f"Error modifying script: {str(e)}")
            return f"Error modifying script: {str(e)}"

# Initialize processors
document_processor = DocumentProcessor()
video_processor = VideoProcessor()
script_generator = EnhancedScriptGenerator()

@app.route('/')
def index():
    return render_template('index.html')
# NEW ROUTES FOR DOCUMENT PROCESSING

@app.route('/api/upload-document', methods=['POST'])
def upload_document():
    """Handle document upload and processing"""
    user_id = session.get('user_id', 'default_user')
    if 'user_id' not in session:
        session['user_id'] = user_id
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    file_id = request.form.get('file_id', str(uuid.uuid4()))
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not document_processor.allowed_file(file.filename):
        return jsonify({'error': 'File type not supported'}), 400
    
    try:
        # Secure filename and save temporarily
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}_{filename}")
        file.save(file_path)
        
        # Process the document
        result = document_processor.process_document(file_path, filename)
        
        # Clean up temp file
        try:
            os.remove(file_path)
        except:
            pass
        
        if result['error']:
            return jsonify({'error': result['error']}), 400
        
        # Store document data
        user_data[user_id]['documents'][file_id] = {
            'filename': filename,
            'text': result['text'],
            'stats': result['stats'],
            'file_type': result['file_type'],
            'processed_at': datetime.now().isoformat(),
            'status': 'completed'
        }
        
        # Clear analysis cache to trigger re-analysis
        user_data[user_id]['analysis_cache'] = {}
        
        logger.info(f"Document processed successfully: {filename} ({result['stats']['char_count']} chars)")
        
        return jsonify({
            'success': True,
            'file_id': file_id,
            'filename': filename,
            'stats': result['stats'],
            'message': 'Document processed successfully'
        })
        
    except Exception as e:
        logger.error(f"Error processing document upload: {str(e)}")
        return jsonify({'error': f'Error processing document: {str(e)}'}), 500

@app.route('/api/documents/<file_id>', methods=['DELETE'])
def delete_document(file_id):
    """Delete a document"""
    user_id = session.get('user_id', 'default_user')
    
    if file_id in user_data[user_id]['documents']:
        del user_data[user_id]['documents'][file_id]
        user_data[user_id]['analysis_cache'] = {}  # Clear cache
        return jsonify({'success': True})
    
    return jsonify({'error': 'Document not found'}), 404

@app.route('/api/documents', methods=['GET'])
def get_documents():
    """Get all uploaded documents"""
    user_id = session.get('user_id', 'default_user')
    documents = user_data[user_id]['documents']
    
    doc_list = []
    for doc_id, doc_data in documents.items():
        doc_list.append({
            'id': doc_id,
            'filename': doc_data['filename'],
            'file_type': doc_data['file_type'],
            'stats': doc_data['stats'],
            'status': doc_data['status'],
            'processed_at': doc_data['processed_at']
        })
    
    return jsonify({'documents': doc_list})

# ENHANCED EXISTING ROUTES

@app.route('/api/analyze-content', methods=['POST'])
def analyze_content():
    """Enhanced analysis including documents"""
    user_id = session.get('user_id', 'default_user')
    folders = user_data[user_id]['folders']
    documents = user_data[user_id]['documents']
    
    # Check if already analyzed
    if user_data[user_id]['analysis_cache'].get('analyzed'):
        return jsonify({
            'success': True,
            'style_profile': user_data[user_id]['analysis_cache'].get('style_profile', ''),
            'inspiration_summary': user_data[user_id]['analysis_cache'].get('inspiration_summary', ''),
            'document_insights': user_data[user_id]['analysis_cache'].get('document_insights', ''),
            'stats': user_data[user_id]['analysis_cache'].get('stats', {}),
            'redirect_to_custom': True
        })
    
    # Gather video content
    personal_transcripts = []
    inspiration_transcripts = []
    
    for folder_data in folders.values():
        for video_data in folder_data['videos'].values():
            if video_data['status'] == 'completed' and video_data.get('transcript'):
                if folder_data['type'] == 'personal':
                    personal_transcripts.append(video_data['transcript'])
                else:
                    inspiration_transcripts.append(video_data['transcript'])
    
    # Gather document content
    document_texts = []
    for doc_data in documents.values():
        if doc_data['status'] == 'completed' and doc_data.get('text'):
            document_texts.append(doc_data['text'])
    
    try:
        # Analyze all content sources
        style_profile = ""
        inspiration_summary = ""
        document_insights = ""
        
        if personal_transcripts:
            logger.info("Starting enhanced style analysis...")
            style_profile = script_generator.analyze_creator_style(personal_transcripts)
        else:
            style_profile = "No personal videos provided. Will use professional YouTube style."
        
        if inspiration_transcripts:
            logger.info("Starting enhanced inspiration analysis...")
            inspiration_summary = script_generator.analyze_inspiration_content(inspiration_transcripts)
        else:
            inspiration_summary = "No inspiration videos provided. Will create original content."
        
        if document_texts:
            logger.info("Starting document analysis...")
            document_insights = script_generator.analyze_documents(document_texts)
        else:
            document_insights = "No documents provided. Will rely on video content and general knowledge."
        
        # Cache results
        user_data[user_id]['analysis_cache'] = {
            'style_profile': style_profile,
            'inspiration_summary': inspiration_summary,
            'document_insights': document_insights,
            'stats': {
                'personal_videos': len(personal_transcripts),
                'inspiration_videos': len(inspiration_transcripts),
                'documents': len(document_texts)
            },
            'timestamp': datetime.now().isoformat(),
            'analyzed': True
        }
        
        logger.info("Enhanced analysis completed successfully")
        
        return jsonify({
            'success': True,
            'style_profile': style_profile,
            'inspiration_summary': inspiration_summary,
            'document_insights': document_insights,
            'stats': {
                'personal_videos': len(personal_transcripts),
                'inspiration_videos': len(inspiration_transcripts),
                'documents': len(document_texts)
            },
            'redirect_to_custom': True
        })
        
    except Exception as e:
        logger.error(f"Error analyzing content: {str(e)}")
        return jsonify({'error': f'Error analyzing content: {str(e)}'}), 500

@app.route('/api/generate-from-prompt', methods=['POST'])
def generate_from_prompt():
    """Enhanced script generation with documents"""
    user_id = session.get('user_id', 'default_user')
    
    data = request.json
    user_prompt = data.get('prompt', '').strip()
    
    if not user_prompt:
        return jsonify({'error': 'Please provide a prompt for script generation'}), 400
    
    # Get or create analysis
    analysis_cache = user_data[user_id]['analysis_cache']
    
    if not analysis_cache.get('analyzed'):
        # Perform quick analysis
        folders = user_data[user_id]['folders']
        documents = user_data[user_id]['documents']
        
        personal_transcripts = []
        inspiration_transcripts = []
        document_texts = []
        
        # Gather video content
        for folder_data in folders.values():
            for video_data in folder_data['videos'].values():
                if video_data['status'] == 'completed' and video_data.get('transcript'):
                    if folder_data['type'] == 'personal':
                        personal_transcripts.append(video_data['transcript'])
                    else:
                        inspiration_transcripts.append(video_data['transcript'])
        
        # Gather document content
        for doc_data in documents.values():
            if doc_data['status'] == 'completed' and doc_data.get('text'):
                document_texts.append(doc_data['text'])
        
        # Quick analysis
        if personal_transcripts:
            style_profile = script_generator.analyze_creator_style(personal_transcripts)
        else:
            style_profile = "Professional, engaging YouTube style with clear explanations and good pacing."
        
        if inspiration_transcripts:
            inspiration_summary = script_generator.analyze_inspiration_content(inspiration_transcripts)
        else:
            inspiration_summary = "Creating original content based on user request and best practices."
        
        if document_texts:
            document_insights = script_generator.analyze_documents(document_texts)
        else:
            document_insights = "No document knowledge provided. Using general expertise and research."
        
        # Cache results
        user_data[user_id]['analysis_cache'] = {
            'style_profile': style_profile,
            'inspiration_summary': inspiration_summary,
            'document_insights': document_insights,
            'stats': {
                'personal_videos': len(personal_transcripts),
                'inspiration_videos': len(inspiration_transcripts),
                'documents': len(document_texts)
            },
            'timestamp': datetime.now().isoformat(),
            'analyzed': True
        }
    else:
        style_profile = analysis_cache.get('style_profile', 'Professional YouTube style')
        inspiration_summary = analysis_cache.get('inspiration_summary', 'Original content creation')
        document_insights = analysis_cache.get('document_insights', 'General knowledge base')
    
    try:
        logger.info(f"Generating enhanced script with documents: {user_prompt[:100]}...")
        
        # Generate script with all knowledge sources
        script = script_generator.generate_enhanced_script(
            style_profile,
            inspiration_summary,
            document_insights,
            user_prompt
        )
        
        # Store current script for chat modifications
        user_data[user_id]['current_script'] = {
            'content': script,
            'style_profile': style_profile,
            'topic_insights': inspiration_summary,
            'document_insights': document_insights,
            'original_prompt': user_prompt,
            'timestamp': datetime.now().isoformat()
        }
        
        # Initialize chat session
        chat_session_id = str(uuid.uuid4())
        user_data[user_id]['chat_sessions'][chat_session_id] = {
            'messages': [],
            'script_versions': [script],
            'created_at': datetime.now().isoformat()
        }
        
        stats = analysis_cache.get('stats', {})
        has_documents = stats.get('documents', 0) > 0
        has_topic_insights = bool(inspiration_summary and inspiration_summary != "Creating original content based on user request and best practices.")
        
        logger.info("Enhanced script generation with documents completed successfully")
        
        return jsonify({
            'success': True,
            'script': script,
            'user_prompt': user_prompt,
            'style_profile': style_profile,
            'inspiration_summary': inspiration_summary,
            'document_insights': document_insights,
            'stats': stats,
            'generation_type': 'enhanced_with_documents' if has_documents else 'analyzed',
            'chat_session_id': chat_session_id,
            'has_topic_insights': has_topic_insights,
            'has_documents': has_documents
        })
        
    except Exception as e:
        logger.error(f"Error generating enhanced script: {str(e)}")
        return jsonify({'error': f'Error generating script: {str(e)}'}), 500

@app.route('/api/chat-modify-script', methods=['POST'])
def chat_modify_script():
    """Enhanced chat modification with document context"""
    user_id = session.get('user_id', 'default_user')
    
    data = request.json
    user_message = data.get('message', '').strip()
    chat_session_id = data.get('chat_session_id')
    
    if not user_message:
        return jsonify({'error': 'Please provide a modification request'}), 400
    
    current_script_data = user_data[user_id].get('current_script')
    if not current_script_data:
        return jsonify({'error': 'No active script to modify'}), 400
    
    try:
        logger.info(f"Processing chat modification with documents: {user_message[:50]}...")
        
        # Get current script and full context
        current_script = current_script_data['content']
        style_profile = current_script_data['style_profile']
        topic_insights = current_script_data['topic_insights']
        document_insights = current_script_data.get('document_insights', '')
        
        # Generate modification with full context
        modification_response = script_generator.modify_script_chat(
            current_script,
            style_profile,
            topic_insights,
            document_insights,
            user_message
        )
        
        # Update chat session
        if chat_session_id and chat_session_id in user_data[user_id]['chat_sessions']:
            chat_session = user_data[user_id]['chat_sessions'][chat_session_id]
            chat_session['messages'].append({
                'user_message': user_message,
                'ai_response': modification_response,
                'timestamp': datetime.now().isoformat()
            })
        
        logger.info("Script modification with documents completed successfully")
        
        return jsonify({
            'success': True,
            'response': modification_response,
            'user_message': user_message,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logger.error(f"Error modifying script via chat: {str(e)}")
        return jsonify({'error': f'Error modifying script: {str(e)}'}), 500

@app.route('/api/status', methods=['GET'])
def get_status():
    """Enhanced status including document info"""
    user_id = session.get('user_id', 'default_user')
    
    status_info = {
        'total_folders': len(user_data[user_id]['folders']),
        'total_videos': 0,
        'completed_videos': 0,
        'processing_videos': 0,
        'error_videos': 0,
        'total_documents': len(user_data[user_id]['documents']),
        'completed_documents': 0,
        'has_personal_videos': False,
        'has_inspiration_videos': False,
        'has_documents': False,
        'ready_for_analysis': False,
        'can_generate_script': False
    }
    
    personal_completed = 0
    inspiration_completed = 0
    
    # Count videos
    for folder_data in user_data[user_id]['folders'].values():
        for video_data in folder_data['videos'].values():
            status_info['total_videos'] += 1
            if video_data['status'] == 'completed':
                status_info['completed_videos'] += 1
                if folder_data['type'] == 'personal':
                    personal_completed += 1
                    status_info['has_personal_videos'] = True
                else:
                    inspiration_completed += 1
                    status_info['has_inspiration_videos'] = True
            elif video_data['status'] == 'processing':
                status_info['processing_videos'] += 1
            elif video_data['status'] == 'error':
                status_info['error_videos'] += 1
    
    # Count documents
    for doc_data in user_data[user_id]['documents'].values():
        if doc_data['status'] == 'completed':
            status_info['completed_documents'] += 1
            status_info['has_documents'] = True
    
    # Determine readiness - can generate with any content source
    has_any_content = (personal_completed > 0 or 
                      inspiration_completed > 0 or 
                      status_info['completed_documents'] > 0)
    
    status_info['ready_for_analysis'] = has_any_content
    status_info['can_generate_script'] = has_any_content
    
    # Analysis quality assessment
    if personal_completed > 0 and inspiration_completed > 0 and status_info['completed_documents'] > 0:
        status_info['analysis_quality'] = 'premium'
    elif (personal_completed > 0 and inspiration_completed > 0) or status_info['completed_documents'] > 0:
        status_info['analysis_quality'] = 'optimal'
    elif personal_completed > 0 or inspiration_completed > 0:
        status_info['analysis_quality'] = 'good'
    else:
        status_info['analysis_quality'] = 'basic'
    
    return jsonify(status_info)

# ALL YOUR EXISTING ROUTES (folders, videos, etc.) - keeping them exactly as they were
@app.route('/api/folders', methods=['GET'])
def get_folders():
    user_id = session.get('user_id', 'default_user')
    folders = user_data[user_id]['folders']
    
    folder_list = []
    for folder_id, folder_data in folders.items():
        folder_list.append({
            'id': folder_id,
            'name': folder_data['name'],
            'type': folder_data['type'],
            'video_count': len(folder_data['videos']),
            'created_at': folder_data['created_at']
        })
    
    return jsonify({'folders': folder_list})

@app.route('/api/folders', methods=['POST'])
def create_folder():
    user_id = session.get('user_id', 'default_user')
    if 'user_id' not in session:
        session['user_id'] = user_id
    
    data = request.json
    folder_name = data.get('name', '').strip()
    folder_type = data.get('type', 'personal')
    
    if not folder_name:
        return jsonify({'error': 'Folder name is required'}), 400
    
    folder_id = str(uuid.uuid4())
    user_data[user_id]['folders'][folder_id] = {
        'name': folder_name,
        'type': folder_type,
        'videos': {},
        'created_at': datetime.now().isoformat()
    }
    
    return jsonify({
        'success': True,
        'folder': {
            'id': folder_id,
            'name': folder_name,
            'type': folder_type,
            'video_count': 0,
            'created_at': user_data[user_id]['folders'][folder_id]['created_at']
        }
    })

@app.route('/api/folders/<folder_id>', methods=['DELETE'])
def delete_folder(folder_id):
    user_id = session.get('user_id', 'default_user')
    
    if folder_id in user_data[user_id]['folders']:
        del user_data[user_id]['folders'][folder_id]
        user_data[user_id]['analysis_cache'] = {}
        return jsonify({'success': True})
    
    return jsonify({'error': 'Folder not found'}), 404

@app.route('/api/folders/<folder_id>/videos', methods=['GET'])
def get_folder_videos(folder_id):
    user_id = session.get('user_id', 'default_user')
    
    if folder_id not in user_data[user_id]['folders']:
        return jsonify({'error': 'Folder not found'}), 404
    
    videos = user_data[user_id]['folders'][folder_id]['videos']
    
    video_list = []
    for video_id, video_data in videos.items():
        video_item = {
            'id': video_id,
            'url': video_data['url'],
            'title': video_data.get('title', 'Unknown Title'),
            'status': video_data['status'],
            'added_at': video_data['added_at'],
            'transcript_length': len(video_data.get('transcript', '')) if video_data.get('transcript') else 0
        }
        
        if video_data.get('stats'):
            video_item['stats'] = video_data['stats']
        
        video_list.append(video_item)
    
    return jsonify({'videos': video_list})

@app.route('/api/folders/<folder_id>/videos', methods=['POST'])
def add_video_to_folder(folder_id):
    user_id = session.get('user_id', 'default_user')
    
    if folder_id not in user_data[user_id]['folders']:
        return jsonify({'error': 'Folder not found'}), 404
    
    data = request.json
    video_url = data.get('url', '').strip()
    
    if not video_url:
        return jsonify({'error': 'Video URL is required'}), 400
    
    if not video_processor.validate_youtube_url(video_url):
        return jsonify({'error': 'Invalid YouTube URL'}), 400
    
    video_id = str(uuid.uuid4())
    user_data[user_id]['folders'][folder_id]['videos'][video_id] = {
        'url': video_url,
        'title': 'Processing...',
        'status': 'pending',
        'transcript': None,
        'stats': None,
        'added_at': datetime.now().isoformat()
    }
    
    user_data[user_id]['analysis_cache'] = {}
    
    def process_video():
        try:
            yt_video_id = video_processor.extract_video_id(video_url)
            title = f"Video {yt_video_id}" if yt_video_id else "Unknown Video"
            
            user_data[user_id]['folders'][folder_id]['videos'][video_id]['status'] = 'processing'
            user_data[user_id]['folders'][folder_id]['videos'][video_id]['title'] = title
            
            result = video_processor.extract_transcript_details(video_url)
            
            if result['error']:
                user_data[user_id]['folders'][folder_id]['videos'][video_id]['status'] = 'error'
                user_data[user_id]['folders'][folder_id]['videos'][video_id]['error'] = result['error']
            else:
                user_data[user_id]['folders'][folder_id]['videos'][video_id]['status'] = 'completed'
                user_data[user_id]['folders'][folder_id]['videos'][video_id]['transcript'] = result['transcript']
                user_data[user_id]['folders'][folder_id]['videos'][video_id]['stats'] = result['stats']
            
        except Exception as e:
            logger.error(f"Error processing video {video_id}: {str(e)}")
            user_data[user_id]['folders'][folder_id]['videos'][video_id]['status'] = 'error'
            user_data[user_id]['folders'][folder_id]['videos'][video_id]['error'] = str(e)
    
    thread = threading.Thread(target=process_video)
    thread.start()
    
    return jsonify({
        'success': True,
        'video': {
            'id': video_id,
            'url': video_url,
            'title': 'Processing...',
            'status': 'pending',
            'added_at': user_data[user_id]['folders'][folder_id]['videos'][video_id]['added_at']
        }
    })

@app.route('/api/folders/<folder_id>/videos/<video_id>', methods=['DELETE'])
def delete_video_from_folder(folder_id, video_id):
    user_id = session.get('user_id', 'default_user')
    
    if folder_id not in user_data[user_id]['folders']:
        return jsonify({'error': 'Folder not found'}), 404
    
    if video_id in user_data[user_id]['folders'][folder_id]['videos']:
        del user_data[user_id]['folders'][folder_id]['videos'][video_id]
        user_data[user_id]['analysis_cache'] = {}
        return jsonify({'success': True})
    
    return jsonify({'error': 'Video not found'}), 404

@app.route('/api/get-topic-insights', methods=['POST'])
def get_topic_insights():
    """Extract specific topic insights for a given subject"""
    user_id = session.get('user_id', 'default_user')
    
    data = request.json
    topic_query = data.get('topic', '').strip()
    
    if not topic_query:
        return jsonify({'error': 'Please provide a topic to analyze'}), 400
    
    # Get inspiration content and documents
    folders = user_data[user_id]['folders']
    documents = user_data[user_id]['documents']
    
    inspiration_transcripts = []
    document_texts = []
    
    # Gather inspiration videos
    for folder_data in folders.values():
        if folder_data['type'] == 'inspiration':
            for video_data in folder_data['videos'].values():
                if video_data['status'] == 'completed' and video_data.get('transcript'):
                    inspiration_transcripts.append(video_data['transcript'])
    
    # Gather documents
    for doc_data in documents.values():
        if doc_data['status'] == 'completed' and doc_data.get('text'):
            document_texts.append(doc_data['text'])
    
    if not inspiration_transcripts and not document_texts:
        return jsonify({'error': 'No content available for topic analysis. Upload documents or add inspiration videos.'}), 400
    
    try:
        # Create enhanced topic-specific analysis prompt
        topic_analysis_prompt = f"""
        You are an expert topic analyst with access to multiple knowledge sources. Analyze the following content to extract comprehensive insights about: "{topic_query}"

        Focus on finding:

        **MAIN POINTS & KEY CONCEPTS:**
        - Core ideas and principles related to {topic_query}
        - Important definitions and explanations
        - Key statistics, data points, and research findings
        - Expert opinions and industry insights from all sources

        **PRACTICAL APPLICATIONS:**
        - How-to steps and actionable advice
        - Tools, resources, and recommendations mentioned
        - Real-world examples and case studies
        - Success stories and best practices

        **AUTHORITATIVE KNOWLEDGE:**
        - Expert methodologies and proven approaches
        - Research-backed strategies and techniques
        - Industry standards and professional insights
        - Credible sources and reference materials

        **COMMON CHALLENGES & SOLUTIONS:**
        - Frequently mentioned problems in this topic area
        - Proven solutions and workarounds
        - Mistakes to avoid and pitfalls
        - Expert tips for overcoming obstacles

        **TRENDING DISCUSSIONS:**
        - Current debates or controversies
        - Emerging trends and new developments
        - Future predictions and implications
        - Different perspectives on the topic

        **CONTENT OPPORTUNITIES:**
        - Subtopics worth exploring in detail
        - Angles not fully covered
        - Questions audiences frequently ask
        - Advanced concepts that could be simplified

        Provide specific, actionable insights that could inform a comprehensive, authoritative video script about {topic_query}.

        **AVAILABLE CONTENT SOURCES:**
        """
        
        all_content = []
        source_count = {'videos': 0, 'documents': 0}
        
        # Add inspiration video content
        if inspiration_transcripts:
            all_content.extend(inspiration_transcripts)
            source_count['videos'] = len(inspiration_transcripts)
            topic_analysis_prompt += f"\n\n**INSPIRATION VIDEO TRANSCRIPTS ({len(inspiration_transcripts)} videos):**\n"
            topic_analysis_prompt += "\n\n---VIDEO SEPARATOR---\n\n".join(inspiration_transcripts[:5])  # Limit to first 5
        
        # Add document content
        if document_texts:
            all_content.extend(document_texts)
            source_count['documents'] = len(document_texts)
            topic_analysis_prompt += f"\n\n**DOCUMENT CONTENT ({len(document_texts)} documents):**\n"
            topic_analysis_prompt += "\n\n---DOCUMENT SEPARATOR---\n\n".join(document_texts)
        
        # Limit content size for processing
        max_chars = 50000
        if len(topic_analysis_prompt) > max_chars:
            # Truncate content while preserving structure
            available_space = max_chars - len(topic_analysis_prompt.split('**AVAILABLE CONTENT SOURCES:**')[0])
            available_space = available_space // 2  # Split between videos and docs
            
            truncated_content = []
            
            if inspiration_transcripts:
                video_content = "\n\n---VIDEO SEPARATOR---\n\n".join(inspiration_transcripts)
                if len(video_content) > available_space:
                    video_content = video_content[:available_space] + "\n\n[...CONTENT TRUNCATED...]"
                truncated_content.append(f"**INSPIRATION VIDEO TRANSCRIPTS:**\n{video_content}")
            
            if document_texts:
                doc_content = "\n\n---DOCUMENT SEPARATOR---\n\n".join(document_texts)
                if len(doc_content) > available_space:
                    doc_content = doc_content[:available_space] + "\n\n[...CONTENT TRUNCATED...]"
                truncated_content.append(f"**DOCUMENT CONTENT:**\n{doc_content}")
            
            topic_analysis_prompt = topic_analysis_prompt.split('**AVAILABLE CONTENT SOURCES:**')[0] + \
                                  "**AVAILABLE CONTENT SOURCES:**\n\n" + "\n\n".join(truncated_content)
        
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = model.generate_content(
            topic_analysis_prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.3,
                max_output_tokens=3000
            )
        )
        
        if response.text:
            # Cache the insights
            user_data[user_id]['insights_cache'][topic_query] = {
                'insights': response.text,
                'source_count': source_count,
                'timestamp': datetime.now().isoformat()
            }
            
            return jsonify({
                'success': True,
                'topic': topic_query,
                'insights': response.text,
                'analyzed_videos': source_count['videos'],
                'analyzed_documents': source_count['documents'],
                'total_sources': source_count['videos'] + source_count['documents']
            })
        else:
            return jsonify({'error': 'Could not analyze topic insights'}), 500
            
    except Exception as e:
        logger.error(f"Error analyzing topic insights: {str(e)}")
        return jsonify({'error': f'Error analyzing topic: {str(e)}'}), 500

@app.route('/api/update-script', methods=['POST'])
def update_script():
    """Update the current working script"""
    user_id = session.get('user_id', 'default_user')
    
    data = request.json
    new_script = data.get('script', '').strip()
    chat_session_id = data.get('chat_session_id')
    
    if not new_script:
        return jsonify({'error': 'Script content is required'}), 400
    
    try:
        # Update current script
        if user_data[user_id].get('current_script'):
            user_data[user_id]['current_script']['content'] = new_script
            user_data[user_id]['current_script']['timestamp'] = datetime.now().isoformat()
        
        # Add to chat session history
        if chat_session_id and chat_session_id in user_data[user_id]['chat_sessions']:
            chat_session = user_data[user_id]['chat_sessions'][chat_session_id]
            chat_session['script_versions'].append(new_script)
        
        return jsonify({
            'success': True,
            'message': 'Script updated successfully'
        })
        
    except Exception as e:
        logger.error(f"Error updating script: {str(e)}")
        return jsonify({'error': f'Error updating script: {str(e)}'}), 500

@app.route('/api/get-chat-history', methods=['GET'])
def get_chat_history():
    """Get chat history for current session"""
    user_id = session.get('user_id', 'default_user')
    chat_session_id = request.args.get('chat_session_id')
    
    if not chat_session_id or chat_session_id not in user_data[user_id]['chat_sessions']:
        return jsonify({'error': 'Chat session not found'}), 404
    
    chat_session = user_data[user_id]['chat_sessions'][chat_session_id]
    
    return jsonify({
        'success': True,
        'messages': chat_session['messages'],
        'script_versions': len(chat_session['script_versions']),
        'created_at': chat_session['created_at']
    })

# Health check and utility endpoints
@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'version': '2.0-enhanced'
    })

@app.route('/api/clear-cache', methods=['POST'])
def clear_cache():
    """Clear analysis cache - useful for testing"""
    user_id = session.get('user_id', 'default_user')
    user_data[user_id]['analysis_cache'] = {}
    user_data[user_id]['insights_cache'] = {}
    
    return jsonify({
        'success': True,
        'message': 'Cache cleared successfully'
    })

if __name__ == '__main__':
    # Install required packages if not present
    required_packages = ['PyPDF2', 'python-docx', 'PyMuPDF']
    
    try:
        import PyPDF2
        import docx
        import fitz
    except ImportError as e:
        print(f"Missing required package: {e}")
        print("Please install with: pip install PyPDF2 python-docx PyMuPDF")
        exit(1)
    
    print("=" * 60)
    print("ADVANCED YOUTUBE SCRIPT GENERATOR BACKEND")
    print("=" * 60)
    print("Features:")
    print("✓ Document processing (PDF, DOC, DOCX, TXT)")
    print("✓ YouTube video transcript analysis")
    print("✓ Advanced style profiling")
    print("✓ Topic insights extraction")
    print("✓ Enhanced script generation")
    print("✓ Chat-based script modification")
    print("✓ Multi-source content analysis")
    print("=" * 60)
    print("Starting server...")
    
    app.run(debug=True, threaded=True, port=5000)